"""Builders for realistic PDF/DOCX/TXT test fixtures.

The files are generated at runtime with the same libraries the application uses
to read them, so the extraction tests exercise the real parsers instead of
hand-rolled byte strings.
"""

from __future__ import annotations

import io

SAMPLE_PARAGRAPHS = [
    "Cloud native machine learning platforms have changed how teams ship models.",
    "Containerized inference services can be deployed with the same pipeline as any "
    "other web application, which lowers the operational barrier considerably.",
    "This document evaluates whether a small instruction tuned model is sufficient "
    "for summarizing internal reports, and concludes that it is for most cases.",
]


def build_pdf_bytes(paragraphs: list[str] | None = None, pages: int = 1) -> bytes:
    import pymupdf

    blocks = paragraphs or SAMPLE_PARAGRAPHS
    document = pymupdf.open()
    for page_index in range(pages):
        page = document.new_page()
        text = f"Page {page_index + 1}\n\n" + "\n\n".join(blocks)
        page.insert_textbox(pymupdf.Rect(50, 50, 545, 750), text, fontsize=11)
    data = document.tobytes()
    document.close()
    return data


def build_docx_bytes(paragraphs: list[str] | None = None, with_table: bool = True) -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Quarterly Platform Review", level=1)
    for paragraph in paragraphs or SAMPLE_PARAGRAPHS:
        document.add_paragraph(paragraph)
    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Average latency"
        table.cell(1, 1).text = "1.8 seconds"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_txt_bytes(paragraphs: list[str] | None = None) -> bytes:
    return "\n\n".join(paragraphs or SAMPLE_PARAGRAPHS).encode("utf-8")


def long_text(repeat: int = 8) -> str:
    return " ".join(" ".join(SAMPLE_PARAGRAPHS) for _ in range(repeat))
