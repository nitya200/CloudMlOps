"""Document text extraction.

Factory pattern: ``DocumentExtractorFactory`` maps a file type to a concrete
``TextExtractor``. Supporting a new format (RTF, HTML, ...) means adding one
class and registering it - no changes to the callers.
"""

from __future__ import annotations

import io
from abc import ABC, abstractmethod
from dataclasses import dataclass

from app.core.exceptions import ExtractionError, UnsupportedFileTypeError
from app.core.logging import get_logger
from app.models.enums import FileType
from app.utils.file_utils import decode_text, get_extension
from app.utils.text_utils import normalize_text, word_count

logger = get_logger(__name__)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    page_count: int | None = None

    @property
    def word_count(self) -> int:
        return word_count(self.text)

    @property
    def character_count(self) -> int:
        return len(self.text)


class TextExtractor(ABC):
    """Strategy for turning raw file bytes into plain text."""

    file_type: FileType

    @abstractmethod
    def _parse(self, data: bytes) -> ExtractionResult: ...

    def extract(self, data: bytes) -> ExtractionResult:
        try:
            result = self._parse(data)
        except ExtractionError:
            raise
        except Exception as exc:  # third-party parsers raise a wide range
            logger.exception("extraction failed", extra={"file_type": str(self.file_type)})
            raise ExtractionError(
                f"Could not read the {str(self.file_type).upper()} file. "
                "It may be corrupted or password protected."
            ) from exc

        text = normalize_text(result.text)
        if not text.strip():
            raise ExtractionError(
                "No readable text was found in the document. "
                "Scanned images require OCR, which this platform does not perform."
            )
        return ExtractionResult(text=text, page_count=result.page_count)


class PdfExtractor(TextExtractor):
    file_type = FileType.PDF

    def _parse(self, data: bytes) -> ExtractionResult:
        import pymupdf  # imported lazily so the module loads without the wheel

        with pymupdf.open(stream=data, filetype="pdf") as document:
            if document.needs_pass:
                raise ExtractionError("This PDF is password protected.")
            pages = [page.get_text("text") for page in document]
            return ExtractionResult(text="\n\n".join(pages), page_count=document.page_count)


class DocxExtractor(TextExtractor):
    file_type = FileType.DOCX

    def _parse(self, data: bytes) -> ExtractionResult:
        import docx

        document = docx.Document(io.BytesIO(data))
        blocks = [para.text for para in document.paragraphs if para.text.strip()]
        # Tables often carry the substance of a report, so include them too.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return ExtractionResult(text="\n".join(blocks))


class TxtExtractor(TextExtractor):
    file_type = FileType.TXT

    def _parse(self, data: bytes) -> ExtractionResult:
        return ExtractionResult(text=decode_text(data))


class DocumentExtractorFactory:
    """Creates the extractor that matches a file type."""

    _registry: dict[FileType, type[TextExtractor]] = {
        FileType.PDF: PdfExtractor,
        FileType.DOCX: DocxExtractor,
        FileType.TXT: TxtExtractor,
    }

    @classmethod
    def register(cls, file_type: FileType, extractor: type[TextExtractor]) -> None:
        cls._registry[file_type] = extractor

    @classmethod
    def supported_types(cls) -> list[str]:
        return sorted(str(file_type) for file_type in cls._registry)

    @classmethod
    def create(cls, file_type: FileType | str) -> TextExtractor:
        try:
            key = FileType(str(file_type).lower())
        except ValueError as exc:
            raise UnsupportedFileTypeError(f"Unsupported file type '{file_type}'.") from exc
        extractor_cls = cls._registry.get(key)
        if extractor_cls is None:
            raise UnsupportedFileTypeError(f"No extractor registered for '{key}'.")
        return extractor_cls()

    @classmethod
    def create_for_filename(cls, filename: str) -> TextExtractor:
        return cls.create(get_extension(filename))


class ExtractionService:
    """Thin facade so callers do not need to know about the factory."""

    def extract(self, file_type: FileType | str, data: bytes) -> ExtractionResult:
        return DocumentExtractorFactory.create(file_type).extract(data)

    def extract_from_filename(self, filename: str, data: bytes) -> ExtractionResult:
        return DocumentExtractorFactory.create_for_filename(filename).extract(data)
